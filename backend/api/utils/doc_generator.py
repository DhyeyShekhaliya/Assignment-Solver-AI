# backend/api/utils/doc_generator.py

from docx import Document
import os

def generate_answer_doc(question_text, answer_text, output_path):
    doc = Document()
    doc.add_heading("Assignment Answers", level=1)
    doc.add_heading("Question", level=2)
    doc.add_paragraph(question_text)

    doc.add_heading("Answer", level=2)
    doc.add_paragraph(answer_text)

    doc.save(output_path)
