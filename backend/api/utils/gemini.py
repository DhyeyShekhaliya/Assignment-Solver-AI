# backend/api/utils/gemini.py

import os
import requests
import json
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

def extract_docx_content(file_path):
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
    return "\n".join(text), doc

def ask_gemini(assignment_text: str, title: str = "Untitled Assignment"):
    headers = {
        "Content-Type": "application/json"
    }

    prompt = f"""
You are an expert assistant that processes assignments consisting of questions, problems, or tasks provided by the user. Your job is to:

1. Carefully read and understand the assignment content.
2. Answer all questions or solve all problems, providing detailed reasoning.
3. If coding is required:
   - Write the code in Python (unless otherwise specified).
   - Simulate or explain the expected output if actual execution isn't possible.
4. Write a concise conclusion summarizing the work done.

Please return your answers in clear, well-structured text format. Use the following layout:

Assignment Title: {title}

Question 1:
[Insert the question text here]
Answer:
[Answer with explanation]
Code (if any):
[Code here]
Expected Output:
[Expected output]
Remarks:
[Any notes or explanation]

...repeat for all questions...

Conclusion:
[Summarize the work or insights]

Now here is the user's assignment content:

\"\"\"{assignment_text}\"\"\"
"""


    data = {
        "contents": [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ]
    }

    response = requests.post(
        f"{GEMINI_URL}?key={GEMINI_API_KEY}",
        headers=headers,
        json=data
    )

    if response.status_code == 200:
        try:
            return response.json()["candidates"][0]["content"]["parts"][0]["text"]
        except Exception as e:
            return f"[ERROR: Unexpected Gemini response format: {str(e)}]"
    else:
        return f"[ERROR: Gemini API returned status {response.status_code}]"

def format_output_to_docx(gemini_response: dict, output_path: str):
    new_doc = Document()
    style = new_doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    new_doc.add_heading(gemini_response.get("assignment_title", "Processed Assignment"), 0)
    new_doc.add_paragraph(f"Name: {gemini_response.get('name', 'N/A')}")
    new_doc.add_paragraph(f"Roll Number: {gemini_response.get('roll_number', 'N/A')}")

    new_doc.add_heading("Questions & Answers", level=1)
    for idx, item in enumerate(gemini_response["questions"], 1):
        new_doc.add_heading(f"Q{idx}: {item['question']}", level=2)
        new_doc.add_paragraph(f"Answer: {item['answer']}")
        if item["code"]:
            new_doc.add_paragraph("Code:")
            new_doc.add_paragraph(item["code"])
        if item["output"]:
            new_doc.add_paragraph("Expected Output:")
            new_doc.add_paragraph(item["output"])
        if item["remarks"]:
            new_doc.add_paragraph(f"Remarks: {item['remarks']}")

    new_doc.add_heading("Conclusion", level=1)
    new_doc.add_paragraph(gemini_response["conclusion"])

    new_doc.save(output_path)

def process_assignment(file_path, output_path="output/Formatted_Assignment.docx"):
    assignment_text, _ = extract_docx_content(file_path)
    title = os.path.basename(file_path).replace(".docx", "").replace("_", " ")

    gemini_result = ask_gemini(assignment_text, title=title)

    if isinstance(gemini_result, str) and gemini_result.startswith("[ERROR"):
        return gemini_result

    format_output_to_docx(gemini_result, output_path)
    return f"[SUCCESS] Formatted assignment saved at: {output_path}"
